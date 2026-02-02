from pathlib import Path
from typing import Iterator

import docx

from el_pipeline.registry import loaders
from el_pipeline.types import Document


@loaders.register("docx")
class DocxLoader:
    """Loads Microsoft Word documents."""

    def load(self, path: str) -> Iterator[Document]:
        doc = docx.Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        yield Document(id=Path(path).stem, text=text, meta={"source": path})

